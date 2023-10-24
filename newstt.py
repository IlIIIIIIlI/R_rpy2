import os
from flask import Flask, render_template, request, send_file
import pandas as pd
from docx import Document
import re
import chardet
from datetime import datetime
from celery_config import app as celery_app

# import locale
# import rpy2.robjects as robjects

# locale.setlocale(locale.LC_TIME, 'zh_CN.utf-8')


class FileProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.setup_r_environment()
        self.process_with_python()

    def setup_r_environment(self):
        import os

        os.environ["R_HOME"] = f"{os.environ['CONDA_PREFIX']}\\Lib\\R"
        import rpy2.robjects as robjects

        # robjects.r('install.packages("tidyverse",repos="https://mirrors.tuna.tsinghua.edu.cn/CRAN/")')
        # robjects.r('BiocManager::install("seqLogo")')
        robjects.r("library(tidyverse)")
        # robjects.r('install.packages("readxl",repos="https://mirrors.tuna.tsinghua.edu.cn/CRAN/")')
        # robjects.r('BiocManager::install("seqLogo")')
        robjects.r("library(readxl)")
        # robjects.r('install.packages("jiebaR",repos="https://mirrors.tuna.tsinghua.edu.cn/CRAN/")')
        # robjects.r('BiocManager::install("seqLogo")')
        robjects.r("library(jiebaR)")
        robjects.r('Sys.setlocale(category = "LC_ALL", locale = "Chinese")')
        robjects.r("remove(list = ls())")
        print(self.file_path.replace("\\", "/"))
        fixed_path = self.file_path.replace("\\", "/")
        robjects.r(f'sample <- read_excel("{fixed_path}")')
        robjects.r("p <- c(nrow(sample):1)   ## 新闻逆序变成正序")
        robjects.r("sample <- sample[p,]")
        robjects.r('names(sample)[1] <- "time"')
        robjects.r('names(sample)[2] <- "event"')
        robjects.r("data <- filter(sample, is.na(event) != T)")
        robjects.r(
            """
        byte <- c()
        for(i in 1 : nrow(data)){
        byte = c(byte, object.size(data$time[i]))
        num = which(byte != 120)
        }

        undefine_1 <- c()
        if(length(num) != 0){
        undefine_1 = data[num, ]
        data = data[- num, ]
        }

        ####
        ## tag用来做最后的总分类排序（时间顺序）
        ##分词设定 标记词性

        engine = worker(type = "tag")

        ## 引入标记删失词

        delete <- read.csv("12_delete_head.csv")
        dic_delete <- c(delete$delete)
        new_user_word(engine, dic_delete)

        tail <- read.csv("13_delete_contain.csv")
        dic_tail <- c(tail$tail)
        new_user_word(engine, dic_tail)

        ## 引入机构层级词典

        country = read.csv("1_country.csv")
        dic_country <- c(country[,1])
        new_user_word(engine, dic_country)

        ## 引入地区层级词典

        institute <- read.csv("2_institute.csv")
        dic_institute <- c(institute$institute)
        new_user_word(engine, dic_institute)

        ## 引入个体机构级词典

        individual <- read.csv("3_individual.csv")
        dic_individual <- c(individual$individual)
        new_user_word(engine, dic_individual)

        ## 引入省层级词典

        province <- read.csv("4_province.csv")
        dic_province <- c(province$province)
        new_user_word(engine, dic_province)

        ## 引入公司层级词典

        company <- read.csv("5_company.csv")
        dic_company <- c(company$company)
        new_user_word(engine, dic_company)

        ## ？

        forbid <- read.csv("7_reply.csv")
        dic_forbid <- c(forbid$forbid)
        new_user_word(engine, dic_forbid)

        ######################################################################构造total列表

        coun <- tibble(name = country$first_class,
                    tag = country$country_tag_1,
                    table = 1)
        ins <- tibble(name = institute$institute,
                    tag = institute$institute_tag,
                    table = 2)
        indi <- tibble(name = individual$individual,
                    tag = individual$individual_tag,
                    table = 3)
        prov <- tibble(name = province$city,
                    tag = province$provincet_tag,
                    table = 4)
        com <- tibble(name = company$company,
                    tag = company$company_tag,
                    table = 5)
        forb <- tibble(name = forbid$forbid,
                    tag = 1,
                    table = 6)

        total <- rbind(coun, ins, indi, prov, com, forb)[-1,]

        ## 引入消息来源层级词典

        resource <- read.csv("11_press.csv")
        dic_resource <- c(resource$resource)
        new_user_word(engine, dic_resource)

        ## 引用无意义字段

        useless <- read.csv("14_useless.csv")
        dic_useless <- c(useless$useless)
        new_user_word(engine, dic_useless)

        ## 引用政府限制词字段

        gov <- read.csv("8_forbid.csv")
        dic_gov <- c(gov$gov)
        new_user_word(engine, dic_gov)

        ## 引入银行，证券公司字段

        bank <- read.csv("6_bank.csv")
        dic_bank <- c(bank$name)
        new_user_word(engine,dic_bank)

        ## 引用疫情字段

        covid <- read.csv("9_covid.csv")
        dic_covid <- c(covid$covid, covid$n)
        new_user_word(engine, dic_covid)

        ###################################################################
        ## 首先删掉一部分标记新闻
        ## 开头tag

        data$delete <- 0

        for(a in 1:nrow(data)){ ## 第a行新闻
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        for(b in title_table$value[1]){         ## 第二层级匹配
            temp =  which(delete$delete == b)
            if(length(temp) != 0){
            data$delete[a] = 1
            }
        }
        }  

        ##结尾tag
        for(a in 1:nrow(data)){ ## 第a行新闻
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        for(b in title_table$value){         ## 第二层级匹配
            temp =  which(tail$tail == b)
            if(length(temp) != 0){
            data$delete[a] = 1
            }
        }
        }

        delete_news <- filter(data, delete == 1)
        data <- filter(data, delete == 0)
        data <- data[,-3]

        ##
        ##### 删除新闻
        #############################################################################
        ## 对于原油、黄金、贱金属、BDI的分类排序

        energy <- read.csv("10_embassy.csv")
        dic_energy <- c(energy$name)
        new_user_word(engine, dic_energy)

        data$energy <- 0

        for(a in 1:nrow(data)){ ## 第a行新闻
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        for(b in title_table$value){         
            temp =  which(energy$name == b)
            if(length(temp) != 0){
            if(data$energy[a] == 0){
                data$energy[a] = energy$tag[temp][1]
            }
            }
        }
        }  


        data_embassy  <- data %>% filter(energy == 1)


        ####
        ### 原油部分与国家匹配


        data_1 <- data_embassy

        content_1 <- tibble(text = character(length = nrow(data_1)))

        for( i in 1:nrow(data_1)){
        ff = paste(data_1$time[i], data_1$event[i])
        gg = str_replace_all(ff, " ", '')  ## 小标题内容
        content_1$text[i] = gg
        }

        data<- data %>% filter(energy == 0) %>% select(-energy)


        ##############################################################################################

        tag <- tibble(   tag_1 = integer(length = nrow(data)),
                        tag_2 = integer(length = nrow(data)),
                        tag_3 = integer(length = nrow(data)),
                        tag_4 = integer(length = nrow(data)),
                        tag_5 = integer(length = nrow(data)),
                        tag_6 = integer(length = nrow(data)),
                        tag_7 = integer(length = nrow(data)),
                        tag_8 = integer(length = nrow(data)),
                        tag_9 = integer(length = nrow(data)),
                        tag_10 = integer(length = nrow(data)),
                        tag_11 = integer(length = nrow(data)),
                        tag_12 = integer(length = nrow(data)),
                        tag_13 = integer(length = nrow(data)),
                        tag_14 = integer(length = nrow(data)),
                        tag_15 = integer(length = nrow(data)),
                        tag_16 = integer(length = nrow(data)),
                        tag_17 = integer(length = nrow(data)),
                        tag_18 = integer(length = nrow(data)),
                        tag_19 = integer(length = nrow(data)),
                        tag_20 = integer(length = nrow(data)),
                        tag_21 = integer(length = nrow(data)),
                        tag_22 = integer(length = nrow(data)),
                        tag_23 = integer(length = nrow(data)),
                        tag_24 = integer(length = nrow(data)),
                        tag_25 = integer(length = nrow(data)),
                        tag_26 = integer(length = nrow(data)),
                        tag_27 = integer(length = nrow(data)),
                        tag_28 = integer(length = nrow(data)),
                        tag_29 = integer(length = nrow(data)),
                        tag_30 = integer(length = nrow(data)),
                        tag_31 = integer(length = nrow(data)),
                        tag_32 = integer(length = nrow(data)),
                        tag_33 = integer(length = nrow(data)),
                        tag_34 = integer(length = nrow(data)),
                        tag_35 = integer(length = nrow(data)),
                        tag_36 = integer(length = nrow(data)),
                        tag_37 = integer(length = nrow(data)),
                        tag_38 = integer(length = nrow(data)),
                        tag_39 = integer(length = nrow(data)),
                        tag_40 = integer(length = nrow(data)))   ## 公司项目容错1项

        ff <- ncol(tag)
        ## 总体匹配过程要写10重嵌套

        for(a in 1:nrow(data)){ ## 第a行新闻
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        col = 1 ## 国家项目起始列数
        q = nrow(title_table)
        for(b in 1:q){         
            temp =  which(total$name == title_table$value[b])
            if(length(temp) != 0){
            tag[a, col] = total$tag[temp][1]
            col = col + 1
            tag[a, col] = total$table[temp][1]
            col = col + 1
            }
        }
        }  

        for(i in 1:nrow(tag)){
        if(tag$tag_2[i] == 6){
            tag$tag_1[i] = 0
            tag$tag_2[i] = 0
        }
        }

        ## 总体匹配过程
        #############################################################

        data$final_country <- c(0)
        data$final_institute <- c(10000)


        ## 对于首先出现的是国家的直接储存为最终国家
        coun_tag =  which(tag$tag_2 == 1)
        for(i in coun_tag){
        data$final_country[i] = tag$tag_1[i]
        }

        data$final_country <- unlist(data$final_country)
        ##  国家的最终机构确定为4,6,8行等先出现大于1的机构
        ## qieyaoo xiaou  6

        cc <- c(2:20)*2


        for(i in coun_tag){
        bb = which(tag[i,cc] > 1)
        b_2 = which(tag[i, cc] < 6)
        bb = intersect(bb, b_2)
        if(length(bb) != 0){
            m = min(cc[bb])
            data$final_institute[i] = tag[i, m - 1]
        }
        }

        ## 对于先出现的机构，直接储存成最终机构
        ins_tag <- which(tag$tag_2 > 1)
        for(i in ins_tag){
        data$final_institute[i] = tag$tag_1[i]
        }

        ## 中国部门中，如果出现了forbid中的词
        ## 以forbid中的词之前的机构为准，着这样的新闻经常会出现在中国外交部、商务部相关新闻中

        nn <- institute$institute_tag[which(institute$institute == "商务部")][1]
        cc <- c(2:20)*2

        for(i in 1:nrow(data)){
        yy = cc[which(tag[i, cc] == 6)]  ## 标记与forbid表中相等的词
        if(length(yy) != 0){
            if(tag[i, min(yy) - 3] == nn){## 外交部
            data$final_institute[i] = tag[i, min(yy) - 3]
            data$final_country[i] = 0
            }   
        }
        }


        ## 为机构添加默认国家

        for(i in which(data$final_country == 0)){
        if(data$final_institute[i] < 100){  ## 独立机构字段
            data$final_country[i] = individual$country_tag_ind[which(individual$individual_tag == data$final_institute[i])][1]
        }
        if(data$final_institute[i] > 100 && data$final_institute[i] < 2000){    ## 独立机构字段
            data$final_country[i] = institute$country_tag_ins[which(institute$institute_tag == data$final_institute[i])][1]
        }
        if(data$final_institute[i] > 2000 && data$final_institute[i] < 3000){   ## 独立机构字段
            data$final_country[i] = individual$country_tag_ind[which(individual$individual_tag == data$final_institute[i])][1]
        }
        if(data$final_institute[i] > 3000 && data$final_institute[i] < 4000){   ## 省市字段
            data$final_country[i] = institute$country_tag_ins[which(institute$institute_tag == data$final_institute[i])][1]
        }
        if(data$final_institute[i] > 5000 && data$final_institute[i] < 6000){    ## 省市字段
            data$final_country[i] = province$country_tag_pro[which(province$provincet_tag == data$final_institute[i])][1]
        }
        
        if(data$final_institute[i] > 8000 && data$final_institute[i] < 10000){  ## 公司字段
            data$final_country[i] = company$country_tag_com[which(company$company_tag == data$final_institute[i])][1]
        }
        }

        data$final_institute = unlist(data$final_institute)

        ########################################################################################
        ####  机构部分的特殊设定

        ## 1、中国各级政府部门与市场消息区别
        ## 在政府新闻中不可以出现研报、预计分析等用词
        ## 省市 不可以出现有限公司字样 but when it contain ogovement is should set as province

        gov_tag <- rep(0, nrow(data))

        for(a in 1:nrow(data)){ ## 第a行新闻
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        for(b in title_table$value){         ## 第一层级匹配
            temp =  which(gov$gov == b)
            if(length(temp) != 0){
            gov_tag[a] = 1
            }
        }
        } 

        bb <- which(gov_tag == 1)

        if(length(bb) != 0){
        for(i in bb){
            if(data$final_institute[i] < 8000){
            data$final_institute[i] = 10000   ## 如果机构的tag在190以内，就将其归为市场消息
            }
        }
        }

        ## 2、美国国会部分
        #########################################
        ## 参议院与众议院要优于国会 民主党 共和党的排序
        """
        )

        robjects.r(
            """

        kk <- unique(institute$institute_tag[c(which(institute$institute == "国会") : which(institute$institute == "共和党"))])

        senate <- institute$institute_tag[which(institute$institute == "参议院")]
        repre <- institute$institute_tag[which(institute$institute == "众议院")]

        a = c(1:20)*2

        for(i in 1:nrow(data)){
        bb = unique(as.integer(a[which(tag[i,a] == 2)]))-1
        cc = tag[i, bb]
        if(length(cc) != 0 ){
            if(max(cc) <= max(kk) && min(cc) >= min(kk)){
            if(any(cc == senate)){
                data$final_institute[i] = senate
            } 
            if(any(cc == repre)){
                data$final_institute[i] = repre
            } 
            }
        }
        }

        ####3、中国国家总理的处理  不单列
        ######################################################################

        premier <- institute$institute_tag[which(institute$institute_name == "总理")][1]
        council <- institute$institute_tag[which(institute$institute_name == "国务院")][1]

        for(i in 1:nrow(data)){
        if(data$final_country[i] == 100){  ## 中国的国家tag
            if(data$final_institute[i] == premier){
            data$final_institute[i] == council
            }
            if(data$final_institute[i] == premier + 1){ ## 副总理tag
            data$final_institute[i] == council
            }
        }
        }

        ##########################################################
        ## 4、新闻中出现多个地区的，归到地区的最后一名

        a = c(1:20)*2

        for(i in 1:nrow(data)){
        if(data$final_institute[i] > 5000 && data$final_institute[i] < 6000){ ## 对于已经被分类到地方政府的新闻来说
            cc = unique(as.integer(tag[i, bb]))
            if(length(cc) > 1 ){
            data$final_institute[i] = 10000
            }
        }
        }

        ## 5、特殊公司部分
        ##########################################################
        ## 首先将有关沪深A股的消息还原为市场消息

        data$final_institute[which(data$final_institute > 8900)] <- 10000



        ## 同时存在超过一条相关新闻的单独列出，如果只存在一条消息，就归类到市场消息

        kk_1 <- which(data$final_institute > 8000)
        kk_2 <- which(data$final_institute < 10000)
        kk <- unique(data$final_institute[intersect(kk_1, kk_2)])
        if(length(kk) != 0){
        for(i in kk){
            p = which(data$final_institute == i)
            if(length(p) < 2){
            data$final_institute[p] = 10000
            }
        }
        }


        ########################################################################
        ## 消息来源字段
        ## 只拆分国家未确定的新闻

        re <- which(data$final_country == 0)
        for(a in re){ 
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        for(b in unique(title_table$value)){         ## 第二层级匹配
            temp =  which(resource$resource == b)
            if(length(temp) != 0){
            data$final_country[a] = resource$country_tag_re[temp][1]
            }
        }
        }  

        ##########################################################################


        ## 能源和黄金分类待定


        ##疫情
        ## 疫情部分的政府部门有时候会涉及到国务院、央行等政府部门

        mm <- which(data$final_institute > 228)
        ll <- which(data$final_institute < 520)
        mn <- intersect(ll, mm)
        kk <- which(data$final_institute == 10000)
        aa <- which(data$final_institute > 5000)
        bb <- which(data$final_institute < 6000)
        cc <- intersect(aa, bb)
        ml <- c(kk, mn, cc)

        for(a in ml){ ## 第a行新闻
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        temp_1 = 0
        temp_2 = 0
        y <- c()
        for(b in title_table$value){ 
            y = c(y, is.na(as.integer(b)))
            pp =  which(covid$covid == b)
            if(length(pp) != 0){
            temp_1 = temp_1 + 1
            }
            c = which(covid$n == b)
            if(length(c) != 0){
            temp_2 = temp_2 + 1
            }
            if(temp_1 != 0 && temp_2 != 0 && any(y == T)){
            data$final_institute[a] = 20000  ### 疫情字段定义为20000
            }
        }
        }

        ##########################################################


        ##  银行、高盛、惠誉通常会对一个国家的政策效果进行评论，
        ## 此时需要把这些消息都归于市场消息中


        for(a in 1:nrow(data)){ ## 第a行新闻
        tt = segment(data$event[a], engine)
        title_table <- enframe(tt)
        for(b in title_table$value[1]){          ## 第一层级匹配
            temp =  which(bank$name == b)
            if(length(temp) != 0){
            data$final_institute[a] = 10000
            }
        }
        } 
        """
        )

        robjects.r(
            """
        ###  对于欧元区国家的市场消息， 国家tag应该设为399 再按照时间顺序进行排序

        a = which(data$final_institute == 10000)
        b = a[which(data$final_country[a] == 300)] ## 娆ф床
        c = a[which(data$final_country[a] == 301)] ## 娆х洘
        tt <- c(b, c)

        if(length(tt) != 0){
        data$final_country[tt] = 399
        }

        #####################################################


        data$final_country[which(data$final_institute == 2900)] = 1600
        data$final_country[which(data$final_institute == 2901)] = 1700
        data$final_country[which(data$final_institute == 2902)] = 1800
        data$final_country[which(data$final_institute == 2904)] = 1900

        ##

        data$tag <- c(1:nrow(data))

        data_3 <- data[1,]

        uni_c <- unique(data$final_country)

        if(min(uni_c) == 0){
        uni_c <- uni_c[order(uni_c)][-1]
        }else(
        uni_c <- uni_c[order(uni_c)]
        )


        for(i in uni_c){
        u_1 = which(data$final_country == i)
        uni_ins  = unique(data$final_institute[u_1])###############################
        uni_ins = uni_ins[order(uni_ins)]
        for(t in uni_ins){
            u_3 = which(data$final_institute[u_1] == t)
            data_3 = rbind(data_3, data[u_1,][u_3,])
        }
        }

        data_3 <- data_3[-1,]

        ###################################################################################
        ###########   修改疫情分类
        aa <- which(data_3$final_institute == 20000)
        data_cov <- data_3[aa, ]
        data_cov$final_country <- 1

        ###########################################
        ##  如果新闻中不存在疫情新闻，直接分类会导致数据出错

        if(length(aa) != 0){
        data_3 <- data_3[-aa, ]
        data_3 <- rbind(data_cov, data_3)
        }
        #########   疫情分类end
        #########################################################################



        ###############################################################


        text_1 <- tibble(content = rep(0, nrow(data_3)))

        for(i in 1:nrow(data_3)){
        temp = paste(data_3$time[i], data_3$event[i])
        text_1$content[i] = str_replace_all(temp, " ", '')
        }
        """
        )

        robjects.r(
            """
        uni_c <- unique(data_3$final_country)
        b_1<- c()
        b_2 <- c()
        for(i in uni_c){
        aa= which(data_3$final_country == i)
        bb = unique(data_3$final_institute[aa])
        cc = which(country$country_tag_1 == i)[1]
        dd = country$first_class[cc]
        for (t in bb) {
            ee = institute$institute_name[which(institute$institute_tag == t)][1]
            if(is.na(ee)){
            ee = ''
            }
            ff = paste(dd, ee)
            gg = str_replace_all(ff, " ", '')  ## 小标题内容
            if(i < 100 ){  ############################################ ## 
            gg = individual$individual_name[which(individual$individual_tag == t)][1]
            if(is.na(gg)){
                gg = dd
            } 
            }
            if(2000 < t && t < 3000){  ############################################ ## 省市字段
            gg = individual$individual_name[which(individual$individual_tag == t)][1]
            }
            if(5000 < t && t < 8000){  ############################################ ## 省市字段
            gg = province$province[which(province$provincet_tag == t)][1]
            }
            if(8000 < t && t < 10000){  ############################################ ## 省市字段
            gg = company$company_name[which(company$company_tag == t)][1]
            }
            hh = which(data_3$final_institute == t)
            ii = intersect(aa, hh)
            kk = min(ii)  ## 小标题插入行数
            b_1 <- c(b_1, gg)
            b_2 <- c(b_2, kk)
        }
        }
        """
        )

        dd = robjects.r(
            """
        name <- tibble(text = b_1)

        content <- tibble(text = character(length = nrow(data_3)))

        for( i in 1:nrow(data_3)){
        ff = paste(data_3$time[i], data_3$event[i])
        gg = str_replace_all(ff, " ", '')  ## 小标题内容
        content$text[i] = gg
        }


        final_content <- tibble(text = character(length = 1))
        for(i in 1:length(b_1)){
        if(i <= c(length(b_1)-1)){
            final_content = rbind(final_content, name[i, ])
            gg = c(b_2[i]:c(b_2[i+1]-1))
            final_content = rbind(final_content, content[gg,])
        }
        if(i == length(b_1)){
            final_content = rbind(final_content, name[i, ])
            gg = c(b_2[i]:nrow(data_3))
            final_content = rbind(final_content, content[gg,])
        }
        }

        final_content$text[which(final_content$text == "美国美联储")] = "美联储"
        if(final_content$text[2] == "全球"){
        final_content$text[2] <- "全球疫情"
        }

        content_1 <- rbind(tibble(text = "能源"), content_1)
        final_content <- rbind(final_content, content_1)
        """
        )

        robjects.r(
            """
        c <- setdiff(data$tag, data_3$tag)

        if(length(c) != 0){
        undefine <- data[c, ]
        unde <- tibble(text = character(length = nrow(undefine)))
        
        for(i in 1:nrow(unde)){
            bb = paste(undefine$time[i], undefine$event[i])
            bb = str_replace_all(bb, " ", "")
            unde$text[i] = bb
        }
        
        if(object.size(undefine_1) != 0){
            kk <- tibble(text = rep(0, nrow(undefine_1)))
            for(i in 1:nrow(undefine_1)){
            gg = paste(undefine_1$time[i], undefine_1$event[i])
            gg = str_replace_all(gg, " ", "")
            kk$text[i] = gg
            }
            unde <- rbind(unde, kk)
        }
        }

        unde_1 <-tibble(text = "未分类新闻:")
        unde <- rbind(unde_1, unde)
        final_content$text[1] <- paste("新闻截止时间",sample$time[1], "-", sample$time[nrow(sample)])
        final_content <- rbind(final_content, unde)


        ###################################################################################

        write.csv(final_content, file = "龟44龟.csv")
        """
        )

    def process_with_python(self):
        with open("./龟44龟.csv", "rb") as f:
            result = chardet.detect(f.read())
        encoding = result["encoding"]
        with open("./龟44龟.csv", "r", encoding=encoding, errors="ignore") as f:
            content = f.read()
        from io import StringIO

        self.df = pd.read_csv(StringIO(content))

    def generate_word(self):
        doc = Document()
        time_pattern = re.compile(r"\d{2}:\d{2}:\d{2}")
        number_pattern = re.compile(r"(?<!\d)(\d{1}\.)(?!\d)")  # 匹配前后都没有其他数字的序号
        # number_pattern = re.compile(r"(?<!\d)(\d+\.)")  # 正则表达式，查找独立的序号
        current_heading = None
        for _, row in self.df.iterrows():
            text = row["text"]
            if not time_pattern.search(text):
                current_heading = doc.add_heading(text, level=2)
            else:
                if current_heading:
                    # 处理段落中的序号
                    processed_text = number_pattern.sub(r"\n\1", text)
                    p = doc.add_paragraph()
                    p.add_run(processed_text)
        today = datetime.now()
        weekdays = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期天"]
        weekday_str = weekdays[today.weekday()]
        formatted_date = today.strftime("%y.%m.%d") + f"（{weekday_str}）"
        self.filename = f"{formatted_date}每日要闻.docx"
        doc.save(self.filename)


from multiprocessing import Pool, current_process

app = Flask(__name__)


def r_code_execution(file_path):
    print("Process ID: ", current_process().pid)
    processor = FileProcessor(file_path)
    processor.generate_word()
    return processor.filename


@app.route("/")
def index():
    return render_template("upload.html")


import uuid
from multiprocessing import Pool


@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "No file part"
    file = request.files["file"]
    if file.filename == "":
        return "No selected file"
    if not file.filename.endswith(".xlsx"):
        return "Invalid file type, please upload an Excel file"

    unique_filename = str(uuid.uuid4()) + ".xlsx"
    file_path = os.path.join("uploads", unique_filename)
    file.save(file_path)

    # 使用进程池来运行R代码
    with Pool() as pool:
        result = pool.apply_async(r_code_execution, (file_path,))
        filename = result.get()  # 等待结果

    # 发送生成的Word文档作为附件
    return send_file(filename, as_attachment=True)


# if __name__ == "__main__":
#     app.run(host="0.0.0.0", port=5000, debug=True, threaded=False)
